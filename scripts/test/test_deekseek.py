import os
from openai import OpenAI
from openai.types.chat import ChatCompletionMessageParam # <--- IMPORT QUAN TRỌNG ĐỂ SỬA LỖI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CẤU HÌNH DEEPSEEK ---
API_KEY = "sk-b2086419ba0a4219b7d322ae0c45db26"  # Thay key của bạn
BASE_URL = "https://api.deepseek.com"

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

# --- DANH SÁCH TỪ KHÓA CẦN LOẠI BỎ ---
BLACKLIST_PHRASES = [
    "[ĐÃ XONG", "[DA XONG", 
    "Tuyệt vời", "Chào bạn", "Dưới đây là", "Sau đây là", 
    "Vâng,", "Ok,", "DeepSeek:", "Tất nhiên", "Dựa trên thông tin",
    "Chắc chắn rồi", "Here is the"
]

# --- HÀM TẠO FILE WORD ---
def create_word_report(content, filename="Bien_Phap_Thi_Cong_DeepSeek.docx"):
    print(f"\n📝 Đang soạn thảo file Word: {filename}...")
    try:
        doc = Document()
        title = doc.add_heading('HỒ SƠ BIỆN PHÁP THI CÔNG', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue 

            is_garbage = False
            for bad_word in BLACKLIST_PHRASES:
                if line.lower().startswith(bad_word.lower()) or "[ĐÃ XONG" in line:
                    is_garbage = True
                    break
            if is_garbage: continue
            
            if line.startswith('# ') or line.upper().startswith('CHƯƠNG') or line.upper().startswith('PHẦN'):
                clean_text = line.replace('#', '').replace('*', '').strip()
                doc.add_heading(clean_text, level=1)
            elif line.startswith('##') or (len(line) > 2 and line[0].isdigit() and line[1] == '.'): 
                clean_text = line.replace('#', '').replace('*', '').strip()
                doc.add_heading(clean_text, level=2)
            elif line.startswith('- ') or line.startswith('* '):
                clean_text = line.replace('- ', '').replace('* ', '').replace('**', '').strip()
                doc.add_paragraph(clean_text, style='List Bullet')
            else:
                clean_text = line.replace('**', '').strip()
                p = doc.add_paragraph(clean_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(filename)
        print(f"✅ XUẤT FILE THÀNH CÔNG! File tại: {os.path.abspath(filename)}")
    except Exception as e:
        print(f"❌ Lỗi khi lưu file Word: {e}")

# --- SYSTEM PROMPT ---
SYSTEM_INSTRUCTION = """
Bạn là Giám đốc Dự án Xây dựng cấp cao. Nhiệm vụ của bạn là soạn thảo "Biện pháp Tổ chức Thi công" (Method Statement) chuyên nghiệp.

TÀI LIỆU THAM KHẢO & TIÊU CHUẨN CẦN TUÂN THỦ (Học từ mẫu Biện pháp T3):
1.  **Hệ thống Tiêu chuẩn:** Luôn trích dẫn các TCVN hiện hành (VD: TCVN 5308-91 về An toàn, TCVN 4453:1995 về BTCT, TCVN 9361:2012 về Nền móng...).
2.  **Cấu trúc văn bản:** Phải chia thành các CHƯƠNG và MỤC rõ ràng (I, II, III...).

QUY TRÌNH XỬ LÝ (REASONING):
Bước 1: Đọc yêu cầu người dùng.
Bước 2: (QUAN TRỌNG) Đánh giá xem đã đủ thông tin đầu vào chưa?
   - Nếu chưa có: Địa điểm, Quy mô, Loại đất nền, Điều kiện mặt bằng, Nguồn điện/nước... -> HỎI LẠI NGƯỜI DÙNG.
   - Tuyệt đối không tự ý "sáng tác" số liệu địa chất hay quy mô nếu không được cung cấp.

Bước 3: Khi đủ thông tin, soạn thảo Kế hoạch theo khung sau:

--- KHUNG BIỆN PHÁP THI CÔNG ---

CHƯƠNG I: CƠ SỞ LẬP PHƯƠNG ÁN
- Liệt kê các căn cứ pháp lý, luật xây dựng, nghị định và TCVN áp dụng.

CHƯƠNG II: GIỚI THIỆU DỰ ÁN
- Tóm tắt mục tiêu, địa điểm, quy mô hạng mục (Móng, Thân, Hoàn thiện, MEP...).

CHƯƠNG III: CÔNG TÁC CHUẨN BỊ
- Tổ chức lán trại, kho bãi (Kho kín, kho hở).
- Kế hoạch huy động vật tư (Xi măng, Cát, Đá, Thép - Yêu cầu kỹ thuật & Bảo quản).
- Điện, nước phục vụ thi công.

CHƯƠNG IV: BIỆN PHÁP KỸ THUẬT CHI TIẾT
1. Công tác Trắc đạc & Định vị công trình.
2. Công tác Đất & San nền (Đào móng, xử lý nước ngầm, đắp đất K95/K98).
3. Công tác Bê tông cốt thép (Ván khuôn, Cốt thép, Đổ & Bảo dưỡng bê tông).
4. Công tác Xây, Trát, Hoàn thiện.
5. Công tác Lắp đặt thiết bị (nếu có: Máy biến áp, Cột thép, Tủ điện...).

CHƯƠNG V: QUẢN LÝ CHẤT LƯỢNG & AN TOÀN
- Quy trình nghiệm thu nội bộ và nghiệm thu với CĐT.
- Biện pháp an toàn lao động, vệ sinh môi trường, phòng chống cháy nổ.

--- KẾT THÚC KHUNG ---

Hãy viết với giọng văn kỹ thuật, dùng các từ ngữ như: "Nhà thầu cam kết", "Tuân thủ nghiêm ngặt", "Nghiệm thu chuyển bước", "Biện pháp thi công cuốn chiếu/dây chuyền".
"""

# --- MAIN FUNCTION ---
def run_deepseek_agent():
    # --- SỬA LỖI TẠI ĐÂY: Khai báo kiểu dữ liệu rõ ràng ---
    messages: list[ChatCompletionMessageParam] = [
        {"role": "system", "content": SYSTEM_INSTRUCTION}
    ]

    print("👷 BẮT ĐẦU HỆ THỐNG LẬP BIỆN PHÁP (DEEPSEEK FIX)")
    print("-" * 60)
    
    user_input = input(">> Bạn (Nhập yêu cầu): ")

    while True:
        if user_input.lower() in ['exit', 'quit']: break
        
        if user_input.lower() in ['in file', 'xuat file', 'word']:
            # Kiểm tra xem tin nhắn cuối có phải của assistant không
            if len(messages) < 2 or messages[-1]['role'] != 'assistant':
                print("⚠️ Chưa có nội dung từ AI để xuất!")
            else:
                # Ép kiểu nội dung về string để tránh lỗi NoneType
                last_content = str(messages[-1].get('content', ''))
                create_word_report(last_content)
            user_input = input("\n>> Bạn: ")
            continue

        # Thêm tin nhắn user vào list
        messages.append({"role": "user", "content": user_input})

        try:
            print("⏳ DeepSeek đang suy nghĩ...")
            
            response = client.chat.completions.create(
                model="deepseek-reasoner", 
                messages=messages, # Bây giờ biến này đã đúng kiểu dữ liệu
                stream=False
            )
            
            ai_content = response.choices[0].message.content
            print(f"\n🤖 DeepSeek:\n{ai_content}\n")
            
            # Thêm câu trả lời AI vào list
            if ai_content:
                messages.append({"role": "assistant", "content": ai_content})
            
            if ai_content and "[ĐÃ XONG - CÓ THỂ XUẤT FILE]" in ai_content:
                print("💡 Đã hoàn thành. Gõ 'word' hoặc 'in file' để lấy file .docx nhé.")
            
            user_input = input(">> Bạn: ")

        except Exception as e:
            print(f"❌ Lỗi kết nối API: {e}")
            break

if __name__ == "__main__":
    run_deepseek_agent()